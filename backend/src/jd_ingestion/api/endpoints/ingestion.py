# Standard library imports
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional

# Third-party imports
from fastapi import APIRouter, BackgroundTasks, Depends, File, HTTPException, UploadFile
from sqlalchemy import func, select, update
from sqlalchemy.ext.asyncio import AsyncSession

# Local imports
from ...config import settings
from ...core.file_discovery import FileDiscovery
from ...processors.content_processor import ProcessedContent
from ...database.connection import get_async_session
from ...database.models import ContentChunk, JobDescription, JobMetadata, JobSection
from ...processors.content_processor import ContentProcessor
from ...utils.circuit_breaker import circuit_breaker_manager
from ...utils.error_handler import handle_errors, retry_on_failure
from ...utils.logging import get_logger

logger = get_logger(__name__)
router = APIRouter()


@router.post("/scan-directory")
@handle_errors(operation_name="scan_directory")
async def scan_directory(
    directory_path: str,
    recursive: bool = True,
    background_tasks: BackgroundTasks = BackgroundTasks(),
):
    """Scan a directory for job description files."""
    try:
        data_path = Path(directory_path)
        if not data_path.exists():
            raise HTTPException(status_code=400, detail="Directory does not exist")

        # Initialize file discovery
        file_discovery = FileDiscovery(data_path)

        # Scan directory
        files_metadata = file_discovery.scan_directory(recursive=recursive)

        # Generate statistics
        stats = file_discovery.get_stats(files_metadata)

        logger.info(
            "Directory scan completed",
            directory=directory_path,
            total_files=len(files_metadata),
            valid_files=stats["valid_files"],
        )

        return {
            "status": "success",
            "directory": directory_path,
            "stats": stats,
            "files": [
                {
                    "file_path": str(f.file_path),
                    "job_number": f.job_number,
                    "classification": f.classification,
                    "language": f.language,
                    "title": f.title,
                    "file_size": f.file_size,
                    "is_valid": f.is_valid,
                    "validation_errors": f.validation_errors,
                }
                for f in files_metadata[:100]  # Limit response size
            ],
            "total_files_found": len(files_metadata),
        }

    except Exception as e:
        logger.error("Directory scan failed", directory=directory_path, error=str(e))
        raise HTTPException(status_code=500, detail=f"Scan failed: {str(e)}")


@router.post("/process-file")
@handle_errors(operation_name="process_single_file")
@retry_on_failure(max_retries=2, base_delay=1.0)
async def process_single_file(
    file_path: str,
    save_to_db: bool = False,
    db: AsyncSession = Depends(get_async_session),
):
    """Process a single job description file and optionally save to database."""
    # Debug endpoint execution
    try:
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            raise HTTPException(status_code=400, detail="File does not exist")

        # Initialize processors
        file_discovery = FileDiscovery(file_path_obj.parent)
        content_processor = ContentProcessor("")

        # Extract file metadata
        file_metadata = file_discovery._extract_file_metadata(file_path_obj)

        # Validate file metadata before processing
        if not file_metadata.is_valid:
            return {
                "status": "error",
                "file_path": file_path,
                "errors": file_metadata.validation_errors,
            }

        # Read and process file content
        try:
            if file_path_obj.suffix.lower() in {".docx", ".doc"}:
                # Handle Word documents with python-docx
                from docx import Document

                doc = Document(file_path_obj)
                raw_content = "\n".join(
                    [paragraph.text for paragraph in doc.paragraphs]
                )
                logger.info(
                    "Successfully read .docx/.doc file",
                    file=str(file_path_obj),
                    content_length=len(raw_content),
                    paragraphs_count=len(doc.paragraphs),
                )
            elif file_path_obj.suffix.lower() == ".pdf":
                # Handle PDF files (for future implementation)
                logger.warning(
                    "PDF processing not yet implemented", file=str(file_path_obj)
                )
                raw_content = f"PDF content extraction not yet implemented for: {file_path_obj.name}"
            else:
                # Handle text files with detected encoding
                with open(file_path_obj, "r", encoding=file_metadata.encoding) as f:
                    raw_content = f.read()
                logger.info(
                    "Successfully read text file",
                    file=str(file_path_obj),
                    encoding=file_metadata.encoding,
                    content_length=len(raw_content),
                )
        except Exception as e:
            logger.error(
                "Failed to read file",
                file=str(file_path_obj),
                extension=file_path_obj.suffix.lower(),
                error=str(e),
            )
            raise HTTPException(
                status_code=500, detail=f"Failed to read file: {str(e)}"
            )

        # Process content
        try:
            logger.info("Starting content processing", content_length=len(raw_content))
            processed_content = content_processor.process_content(
                raw_content, file_metadata.language or "en"
            )
            logger.info(
                "Content processing succeeded",
                sections_count=len(processed_content.sections),
                has_structured_fields=bool(processed_content.structured_fields),
                sections_keys=(
                    list(processed_content.sections.keys())
                    if processed_content.sections
                    else []
                ),
                structured_fields_attrs=(
                    dir(processed_content.structured_fields)
                    if processed_content.structured_fields
                    else []
                ),
            )
        except Exception as e:
            logger.error("Content processing failed", error=str(e))
            # Create minimal processed_content object for fallback

            @dataclass
            class MinimalStructuredFields:
                position_title: str = ""
                job_number: str = ""
                classification: str = ""
                department: str = ""
                reports_to: str = ""

            @dataclass
            class MinimalProcessedContent:
                cleaned_content: str = raw_content
                sections: Dict = field(default_factory=dict)
                structured_fields: MinimalStructuredFields = field(
                    default_factory=MinimalStructuredFields
                )
                processing_errors: list = field(default_factory=list)

            processed_content = ProcessedContent()
            processed_content.processing_errors = [str(e)]

        # Generate chunks
        try:
            # Ensure cleaned_content is a string before passing to chunk_content
            cleaned_content = processed_content.cleaned_content
            logger.info(
                "About to generate chunks",
                cleaned_content_type=type(cleaned_content).__name__,
            )
            if not isinstance(cleaned_content, str):
                logger.error(
                    "cleaned_content is not a string",
                    content_type=type(cleaned_content).__name__,
                    content_preview=str(cleaned_content)[:200],
                )
                cleaned_content = str(cleaned_content)

            chunks = content_processor.chunk_content(cleaned_content)
            logger.info(
                "Chunk generation completed successfully", chunk_count=len(chunks)
            )
        except Exception as e:
            logger.error("Chunk generation failed", error=str(e))
            chunks = []

        logger.info("About to save to database", save_to_db=save_to_db)

        job_id = None
        if save_to_db:
            try:
                logger.info("Starting database save operation")
                # Save to database

                # Generate unique job number if not available
                job_number = file_metadata.job_number
                if not job_number:
                    # Generate unique job number from file hash and timestamp
                    import hashlib
                    import time

                    timestamp = str(int(time.time()))[-6:]  # Last 6 digits of timestamp
                    file_hash_short = (
                        file_metadata.file_hash[:6].upper()
                        if file_metadata.file_hash
                        else hashlib.md5(file_path_obj.name.encode())
                        .hexdigest()[:6]
                        .upper()
                    )
                    job_number = f"{file_hash_short}-{timestamp}"
                    logger.info(
                        "Generated unique job number",
                        job_number=job_number,
                        filename=file_path_obj.name,
                    )

                # Validate metadata and collect warnings
                metadata_warnings = []
                if not file_metadata.title:
                    metadata_warnings.append("No title extracted from filename - using 'Untitled'")
                    logger.warning("Missing title in file", filename=file_path_obj.name)

                if not file_metadata.classification:
                    metadata_warnings.append("No classification extracted from filename - using 'UNKNOWN'")
                    logger.warning("Missing classification in file", filename=file_path_obj.name)

                # Log all validation warnings
                if metadata_warnings:
                    logger.warning(
                        "File has metadata quality issues",
                        filename=file_path_obj.name,
                        warnings=metadata_warnings
                    )

                # Create job description record with minimal safe data
                job_description = JobDescription(
                    job_number=job_number,
                    title=file_metadata.title or "Untitled",
                    classification=file_metadata.classification or "UNKNOWN",
                    language=file_metadata.language or "en",
                    file_path=str(file_path_obj),
                    raw_content=raw_content,
                    file_hash=file_metadata.file_hash,
                )

                db.add(job_description)
                await db.commit()
                job_id = job_description.id

                logger.info("Job saved to database", job_id=job_id, file_path=file_path)

                # Debug processed_content before saving sections/metadata
                logger.info(
                    "Debug processed_content",
                    has_processed_content=bool(processed_content),
                    has_sections_attr=(
                        hasattr(processed_content, "sections")
                        if processed_content
                        else False
                    ),
                    sections_content=(
                        processed_content.sections
                        if processed_content and hasattr(processed_content, "sections")
                        else None
                    ),
                    has_structured_fields_attr=(
                        hasattr(processed_content, "structured_fields")
                        if processed_content
                        else False
                    ),
                    structured_fields_content=(
                        processed_content.structured_fields
                        if processed_content
                        and hasattr(processed_content, "structured_fields")
                        else None
                    ),
                )

                # Save job sections if processing succeeded
                if (
                    processed_content
                    and hasattr(processed_content, "sections")
                    and processed_content.sections
                ):
                    section_order = 1
                    for (
                        section_type,
                        section_content,
                    ) in processed_content.sections.items():
                        job_section = JobSection(
                            job_id=job_id,
                            section_type=section_type,
                            section_content=section_content,
                            section_order=section_order,
                        )
                        db.add(job_section)
                        section_order += 1

                    logger.info(
                        f"Saved {len(processed_content.sections)} sections to database",
                        job_id=job_id,
                    )

                # Save job metadata if available
                if (
                    processed_content
                    and hasattr(processed_content, "structured_fields")
                    and processed_content.structured_fields
                ):
                    fields = processed_content.structured_fields
                    job_metadata = JobMetadata(
                        job_id=job_id,
                        reports_to=fields.reports_to,
                        department=fields.department,
                        location=fields.location,
                        fte_count=fields.fte_count,
                        salary_budget=fields.salary_budget,
                    )
                    db.add(job_metadata)
                    logger.info("Saved job metadata to database", job_id=job_id)

                # Save content chunks if available
                if chunks:
                    for i, chunk_text in enumerate(chunks):
                        content_chunk = ContentChunk(
                            job_id=job_id, chunk_text=chunk_text, chunk_index=i
                        )
                        db.add(content_chunk)

                    logger.info(
                        f"Saved {len(chunks)} content chunks to database", job_id=job_id
                    )

                # Commit the basic data first
                await db.commit()

                # Generate embeddings for the content chunks (done after commit to have chunk IDs)
                if chunks:
                    try:
                        from ...services.embedding_service import embedding_service

                        logger.info(
                            "Starting embedding generation for content chunks",
                            chunk_count=len(chunks),
                        )

                        # Fetch the saved chunks with their IDs
                        saved_chunks_query = select(ContentChunk).where(
                            ContentChunk.job_id == job_id
                        )
                        saved_chunks_result = await db.execute(saved_chunks_query)
                        saved_chunks = saved_chunks_result.scalars().all()

                        # Generate embeddings in batches
                        chunk_texts = [chunk.chunk_text for chunk in saved_chunks]
                        embeddings = await embedding_service.generate_embeddings_batch(
                            chunk_texts, batch_size=50
                        )

                        # Update chunks with embeddings
                        embedding_count = 0
                        for chunk, embedding in zip(saved_chunks, embeddings):
                            if (
                                embedding
                            ):  # Only update if embedding generation succeeded
                                await db.execute(
                                    update(ContentChunk)
                                    .where(ContentChunk.id == chunk.id)
                                    .values(embedding=embedding)
                                )
                                embedding_count += 1

                        await db.commit()
                        logger.info(
                            f"Generated and saved {embedding_count} embeddings out of {len(chunks)} chunks",
                            job_id=job_id,
                        )

                    except Exception as embedding_error:
                        logger.error(
                            "Embedding generation failed",
                            job_id=job_id,
                            error=str(embedding_error),
                        )
                        # Don't fail the entire process if embeddings fail
                        # The main data has already been committed, no need to rollback

                # Extract skills from job description (done after commit to have job ID)
                try:
                    from ...services.skill_extraction_service import (
                        skill_extraction_service,
                    )

                    logger.info(
                        "Starting skill extraction from job description",
                        job_id=job_id,
                    )

                    # Extract skills from the raw content
                    extracted_skills = await skill_extraction_service.extract_and_save_skills(
                        job_id=job_id,
                        job_text=raw_content,
                        db=db,
                        confidence_threshold=0.5,  # Only include skills with 50%+ confidence
                    )

                    logger.info(
                        f"Extracted and saved {len(extracted_skills)} skills",
                        job_id=job_id,
                        skill_count=len(extracted_skills),
                    )

                except Exception as skill_error:
                    logger.error(
                        "Skill extraction failed",
                        job_id=job_id,
                        error=str(skill_error),
                    )
                    # Don't fail the entire process if skill extraction fails
                    # The main data has already been committed

            except Exception as e:
                logger.error("Database save failed", error=str(e))
                await db.rollback()
                # Don't fail the request, just continue without saving

        logger.info("About to generate response")

        # Safe function for response serialization
        def safe_response_str(value):
            logger.info(
                "safe_response_str called",
                value_type=type(value).__name__,
                value_preview=str(value)[:100],
            )
            if value is None:
                return None
            elif isinstance(value, str):
                return value
            elif isinstance(value, dict):
                # Handle dictionary objects by returning a string representation or None
                logger.warning(
                    "Dictionary passed to safe_response_str", dict_content=value
                )
                return str(value) if value else None
            else:
                # Convert other types to string safely
                try:
                    return str(value)
                except Exception as e:
                    logger.error(
                        "Failed to convert value to string",
                        value_type=type(value).__name__,
                        error=str(e),
                    )
                    return None

        logger.info("Building response dictionary")

        try:
            return {
                "status": "success",
                "file_path": file_path,
                "job_id": job_id,
                "saved_to_database": save_to_db and job_id is not None,
                "warnings": metadata_warnings if metadata_warnings else None,
                "metadata": {
                    "job_number": file_metadata.job_number,
                    "classification": file_metadata.classification,
                    "language": file_metadata.language,
                    "title": file_metadata.title,
                    "file_size": file_metadata.file_size,
                    "file_hash": file_metadata.file_hash,
                },
                "processed_content": {
                    "sections_found": len(processed_content.sections),
                    "sections": list(processed_content.sections.keys()),
                    "structured_fields": {
                        "position_title": safe_response_str(
                            processed_content.structured_fields.position_title
                        ),
                        "job_number": safe_response_str(
                            processed_content.structured_fields.job_number
                        ),
                        "classification": safe_response_str(
                            processed_content.structured_fields.classification
                        ),
                        "department": safe_response_str(
                            processed_content.structured_fields.department
                        ),
                        "reports_to": safe_response_str(
                            processed_content.structured_fields.reports_to
                        ),
                    },
                    "chunks_generated": len(chunks),
                    "processing_errors": processed_content.processing_errors,
                },
            }
        except Exception as e:
            logger.error("Error building response dictionary", error=str(e))
            raise e

    except Exception as e:
        await db.rollback()
        logger.error("File processing failed", file_path=file_path, error=str(e))
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")


async def process_file_in_background(file_path: str, db: AsyncSession):
    try:
        await process_single_file(file_path, save_to_db=True, db=db)
    except Exception as e:
        logger.error(f"Failed to process file in background: {file_path}", error=str(e))


@router.post("/batch-ingest")
@handle_errors(operation_name="batch_ingest_directory")
@retry_on_failure(max_retries=2, base_delay=2.0)
async def batch_ingest_directory(
    directory_path: str,
    recursive: bool = True,
    max_files: Optional[int] = None,
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: AsyncSession = Depends(get_async_session),
):
    """Start batch ingestion of job description files from a directory."""
    try:
        data_path = Path(directory_path)
        if not data_path.exists():
            raise HTTPException(status_code=400, detail="Directory does not exist")

        # Scan directory first
        file_discovery = FileDiscovery(data_path)
        files_metadata = file_discovery.scan_directory(recursive=recursive)

        # Filter valid files
        valid_files = [f for f in files_metadata if f.is_valid]

        # Limit files if specified
        if max_files:
            valid_files = valid_files[:max_files]

        if not valid_files:
            raise HTTPException(
                status_code=400, detail="No valid files found for processing"
            )

        for file_metadata in valid_files:
            background_tasks.add_task(
                process_file_in_background, str(file_metadata.file_path), db
            )

        logger.info(
            "Batch ingestion started",
            directory=directory_path,
            total_files=len(valid_files),
        )

        return {
            "status": "started",
            "directory": directory_path,
            "files_to_process": len(valid_files),
            "message": "Batch ingestion started. Use /jobs/status endpoint to track progress.",
        }

    except Exception as e:
        logger.error("Batch ingestion failed", directory=directory_path, error=str(e))
        raise HTTPException(status_code=500, detail=f"Batch ingestion failed: {str(e)}")


@router.post("/upload")
@handle_errors(operation_name="upload_file")
@retry_on_failure(max_retries=2, base_delay=1.0)
async def upload_file(
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: AsyncSession = Depends(get_async_session),
):
    """Upload and process a single job description file."""
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")

        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in settings.supported_extensions_list:
            raise HTTPException(
                status_code=400, detail=f"Unsupported file extension: {file_ext}"
            )

        # Check file size
        if (
            hasattr(file, "size")
            and file.size > settings.max_file_size_mb * 1024 * 1024
        ):
            raise HTTPException(
                status_code=400,
                detail=f"File too large. Maximum size: {settings.max_file_size_mb}MB",
            )

        # Save uploaded file temporarily
        upload_dir = settings.data_path / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)

        temp_file_path = upload_dir / file.filename

        try:
            with open(temp_file_path, "wb") as f:
                content = await file.read()
                f.write(content)

            # Process the uploaded file and save to database
            result = await process_single_file(
                str(temp_file_path), save_to_db=True, db=db
            )

            return {
                "status": "success",
                "filename": file.filename,
                "processing_result": result,
            }

        finally:
            # Always clean up temporary file
            if temp_file_path.exists():
                temp_file_path.unlink()

    except Exception as e:
        logger.error("File upload failed", filename=file.filename, error=str(e))
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


@router.get("/stats")
@handle_errors(operation_name="get_ingestion_stats")
async def get_ingestion_stats(db: AsyncSession = Depends(get_async_session)):
    """Get comprehensive ingestion statistics."""
    try:
        # Get total job count
        total_result = await db.execute(
            select(func.count()).select_from(JobDescription)
        )
        total_jobs = total_result.scalar_one()

        # Get jobs by classification
        classification_result = await db.execute(
            select(JobDescription.classification, func.count()).group_by(
                JobDescription.classification
            )
        )
        by_classification = {row[0]: row[1] for row in classification_result.fetchall()}

        # Get jobs by language
        language_result = await db.execute(
            select(JobDescription.language, func.count()).group_by(
                JobDescription.language
            )
        )
        by_language = {row[0]: row[1] for row in language_result.fetchall()}

        # Get embedding statistics
        total_chunks_result = await db.execute(
            select(func.count()).select_from(ContentChunk)
        )
        total_chunks = total_chunks_result.scalar_one()

        embedded_chunks_result = await db.execute(
            select(func.count())
            .select_from(ContentChunk)
            .where(ContentChunk.embedding.isnot(None))
        )
        embedded_chunks = embedded_chunks_result.scalar_one()

        # Get jobs with embedding coverage
        jobs_with_embeddings_result = await db.execute(
            select(func.count(func.distinct(ContentChunk.job_id)))
            .select_from(ContentChunk)
            .where(ContentChunk.embedding.isnot(None))
        )
        jobs_with_embeddings = jobs_with_embeddings_result.scalar_one()

        # Get content processing statistics
        jobs_with_sections_result = await db.execute(
            select(func.count(func.distinct(JobSection.job_id))).select_from(JobSection)
        )
        jobs_with_sections = jobs_with_sections_result.scalar_one()

        jobs_with_metadata_result = await db.execute(
            select(func.count(func.distinct(JobMetadata.job_id))).select_from(
                JobMetadata
            )
        )
        jobs_with_metadata = jobs_with_metadata_result.scalar_one()

        # Calculate processing quality metrics
        processing_quality = {
            "jobs_with_sections": jobs_with_sections,
            "jobs_with_metadata": jobs_with_metadata,
            "jobs_with_embeddings": jobs_with_embeddings,
            "section_coverage_rate": (
                (jobs_with_sections / total_jobs * 100) if total_jobs > 0 else 0
            ),
            "metadata_coverage_rate": (
                (jobs_with_metadata / total_jobs * 100) if total_jobs > 0 else 0
            ),
            "embedding_coverage_rate": (
                (jobs_with_embeddings / total_jobs * 100) if total_jobs > 0 else 0
            ),
        }

        # Get section type distribution
        section_types_result = await db.execute(
            select(JobSection.section_type, func.count()).group_by(
                JobSection.section_type
            )
        )
        section_distribution = {
            row[0]: row[1] for row in section_types_result.fetchall()
        }

        # Get recent activity (last 7 days)
        seven_days_ago = datetime.now() - timedelta(days=7)
        recent_jobs_result = await db.execute(
            select(func.count())
            .select_from(JobDescription)
            .where(JobDescription.created_at >= seven_days_ago)
        )
        recent_jobs = recent_jobs_result.scalar_one()

        # Get last updated time
        last_updated_result = await db.execute(
            select(func.max(JobDescription.updated_at))
        )
        last_updated = last_updated_result.scalar_one_or_none()

        # Enhanced processing status based on actual data quality
        processing_status = {
            "completed": jobs_with_embeddings,  # Jobs fully processed with embeddings
            "partial": total_jobs - jobs_with_embeddings,  # Jobs without embeddings
            "needs_embeddings": total_jobs - jobs_with_embeddings,
            "needs_sections": total_jobs - jobs_with_sections,
            "needs_metadata": total_jobs - jobs_with_metadata,
        }

        stats = {
            "total_jobs": total_jobs,
            "by_classification": by_classification,
            "by_language": by_language,
            "processing_status": processing_status,
            "embedding_stats": {
                "total_chunks": total_chunks,
                "embedded_chunks": embedded_chunks,
                "embedding_completion_rate": (
                    (embedded_chunks / total_chunks * 100) if total_chunks > 0 else 0
                ),
                "jobs_with_embeddings": jobs_with_embeddings,
            },
            "content_quality": processing_quality,
            "section_distribution": section_distribution,
            "recent_activity": {
                "jobs_last_7_days": recent_jobs,
                "daily_average": recent_jobs / 7 if recent_jobs > 0 else 0,
            },
            "last_updated": last_updated.isoformat() if last_updated else None,
        }

        return stats

    except Exception as e:
        logger.error("Failed to get ingestion stats", error=str(e))
        raise HTTPException(
            status_code=500, detail="Failed to retrieve ingestion statistics"
        )


@router.get("/task-stats")
@handle_errors(operation_name="get_task_stats")
async def get_task_stats():
    """Get background task processing statistics."""
    try:
        from ...tasks.celery_app import celery_app

        # Get active tasks
        inspect = celery_app.control.inspect()

        # Get task statistics
        stats_dict = {
            "active_tasks": 0,
            "scheduled_tasks": 0,
            "reserved_tasks": 0,
            "workers_online": 0,
            "queue_stats": {
                "processing": {"active": 0, "reserved": 0},
                "embeddings": {"active": 0, "reserved": 0},
            },
            "task_types": {},
        }

        try:
            # Get active tasks
            active = inspect.active()
            if active:
                for worker, tasks in active.items():
                    stats_dict["workers_online"] += 1
                    stats_dict["active_tasks"] += len(tasks)

                    for task in tasks:
                        # Count by queue
                        queue = task.get("delivery_info", {}).get(
                            "routing_key", "unknown"
                        )
                        if queue in stats_dict["queue_stats"]:
                            stats_dict["queue_stats"][queue]["active"] += 1

                        # Count by task type
                        task_name = task.get("name", "unknown")
                        stats_dict["task_types"][task_name] = (
                            stats_dict["task_types"].get(task_name, 0) + 1
                        )

            # Get reserved tasks
            reserved = inspect.reserved()
            if reserved:
                for worker, tasks in reserved.items():
                    stats_dict["reserved_tasks"] += len(tasks)

                    for task in tasks:
                        # Count by queue
                        queue = task.get("delivery_info", {}).get(
                            "routing_key", "unknown"
                        )
                        if queue in stats_dict["queue_stats"]:
                            stats_dict["queue_stats"][queue]["reserved"] += 1

            # Get scheduled tasks
            scheduled = inspect.scheduled()
            if scheduled:
                for worker, tasks in scheduled.items():
                    stats_dict["scheduled_tasks"] += len(tasks)

        except Exception as e:
            logger.warning("Failed to get detailed task stats", error=str(e))
            # Return basic stats structure even if worker inspection fails

        return {
            "status": "success",
            "task_stats": stats_dict,
            "timestamp": datetime.now().isoformat(),
        }

    except Exception as e:
        logger.error("Failed to get task stats", error=str(e))
        # Return empty stats instead of error to avoid breaking the API
        return {
            "status": "unavailable",
            "task_stats": {
                "active_tasks": 0,
                "scheduled_tasks": 0,
                "reserved_tasks": 0,
                "workers_online": 0,
                "queue_stats": {
                    "processing": {"active": 0, "reserved": 0},
                    "embeddings": {"active": 0, "reserved": 0},
                },
                "task_types": {},
            },
            "timestamp": datetime.now().isoformat(),
            "error": "Task monitoring unavailable",
        }
        raise HTTPException(status_code=500, detail="Failed to retrieve statistics")


@router.post("/generate-embeddings")
async def generate_embeddings_for_existing_jobs(
    background_tasks: BackgroundTasks = BackgroundTasks(),
    force_regenerate: bool = False,
    job_ids: Optional[List[int]] = None,
    db: AsyncSession = Depends(get_async_session),
):
    """Generate embeddings for existing jobs that don't have embeddings."""
    try:
        # Build query to find chunks without embeddings
        query = select(ContentChunk)

        if not force_regenerate:
            query = query.where(ContentChunk.embedding.is_(None))

        if job_ids:
            query = query.where(ContentChunk.job_id.in_(job_ids))

        result = await db.execute(query)
        chunks_without_embeddings = result.scalars().all()

        if not chunks_without_embeddings:
            return {
                "status": "success",
                "message": "All chunks already have embeddings or no chunks found",
                "chunks_processed": 0,
            }

        logger.info(
            f"Found {len(chunks_without_embeddings)} chunks to process for embeddings"
        )

        # Process embeddings in the background
        background_tasks.add_task(
            _generate_embeddings_background_task,
            chunks_without_embeddings,
            force_regenerate,
            db,
        )

        return {
            "status": "started",
            "message": f"Embedding generation started for {len(chunks_without_embeddings)} chunks",
            "chunks_to_process": len(chunks_without_embeddings),
            "force_regenerate": force_regenerate,
        }

    except Exception as e:
        logger.error("Failed to start embedding generation", error=str(e))
        raise HTTPException(
            status_code=500, detail=f"Failed to start embedding generation: {str(e)}"
        )


async def _generate_embeddings_background_task(
    chunks: List[ContentChunk], force_regenerate: bool, db: AsyncSession
):
    """Background task to generate embeddings for content chunks."""
    try:
        from ...services.embedding_service import embedding_service

        logger.info(
            f"Starting background embedding generation for {len(chunks)} chunks"
        )

        # Generate embeddings in batches
        batch_size = 50  # Process in smaller batches for better error handling
        total_processed = 0
        total_successful = 0

        for i in range(0, len(chunks), batch_size):
            batch = chunks[i : i + batch_size]
            batch_texts = [chunk.chunk_text for chunk in batch]

            logger.info(
                f"Processing batch {i // batch_size + 1} of {(len(chunks) + batch_size - 1) // batch_size}"
            )

            try:
                # Generate embeddings for this batch
                embeddings = await embedding_service.generate_embeddings_batch(
                    batch_texts, batch_size=batch_size
                )

                # Update chunks with embeddings
                for chunk, embedding in zip(batch, embeddings):
                    if embedding:  # Only update if embedding generation succeeded
                        await db.execute(
                            update(ContentChunk)
                            .where(ContentChunk.id == chunk.id)
                            .values(embedding=embedding)
                        )
                        total_successful += 1

                    total_processed += 1

                await db.commit()
                logger.info(
                    f"Batch {i // batch_size + 1} completed: {sum(1 for e in embeddings if e)} embeddings generated"
                )

            except Exception as batch_error:
                logger.error(
                    f"Batch {i // batch_size + 1} failed", error=str(batch_error)
                )
                await db.rollback()
                # Continue with next batch even if this one fails

        logger.info(
            f"Background embedding generation completed: {total_successful}/{total_processed} chunks processed successfully"
        )

    except Exception as e:
        logger.error("Background embedding generation failed", error=str(e))
        await db.rollback()


@router.get("/resilience-status")
async def get_resilience_status():
    """Get system resilience and circuit breaker status."""
    try:
        # Get circuit breaker metrics
        circuit_metrics = circuit_breaker_manager.get_all_metrics()

        # Get system health indicators
        health_indicators = {
            "embedding_service": {
                "status": (
                    "healthy"
                    if "openai_api" not in circuit_metrics
                    or circuit_metrics["openai_api"]["state"] == "closed"
                    else "degraded"
                ),
                "circuit_breaker": circuit_metrics.get("openai_api", {}),
            },
            "database": {
                "status": (
                    "healthy"
                    if "database" not in circuit_metrics
                    or circuit_metrics["database"]["state"] == "closed"
                    else "degraded"
                ),
                "circuit_breaker": circuit_metrics.get("database", {}),
            },
        }

        # Calculate overall system health
        degraded_services = [
            service
            for service, info in health_indicators.items()
            if info["status"] == "degraded"
        ]

        overall_status = "healthy" if not degraded_services else "degraded"

        return {
            "status": "success",
            "overall_health": overall_status,
            "degraded_services": degraded_services,
            "circuit_breakers": circuit_metrics,
            "health_indicators": health_indicators,
            "timestamp": datetime.now().isoformat(),
            "recommendations": _get_health_recommendations(health_indicators),
        }

    except Exception as e:
        logger.error("Failed to get resilience status", error=str(e))
        return {
            "status": "error",
            "error": str(e),
            "timestamp": datetime.now().isoformat(),
        }


def _get_health_recommendations(health_indicators: Dict[str, Any]) -> List[str]:
    """Generate health recommendations based on current system state."""
    recommendations = []

    for service, info in health_indicators.items():
        if info["status"] == "degraded":
            circuit_info = info.get("circuit_breaker", {})
            state = circuit_info.get("state", "unknown")

            if state == "open":
                recommendations.append(
                    f"{service} service is currently unavailable. "
                    f"Circuit breaker will retry in {circuit_info.get('recovery_timeout', 'unknown')} seconds."
                )
            elif state == "half_open":
                recommendations.append(
                    f"{service} service is being tested for recovery. "
                    "Monitor for successful operations."
                )

            failure_rate = circuit_info.get("failure_rate", 0)
            if failure_rate > 0.5:
                recommendations.append(
                    f"{service} service has high failure rate ({failure_rate:.2%}). "
                    "Consider investigating underlying issues."
                )

    if not recommendations:
        recommendations.append("All services are operating normally.")

    return recommendations


@router.post("/reset-circuit-breakers")
async def reset_circuit_breakers(service_name: Optional[str] = None):
    """Reset circuit breakers to allow recovery testing."""
    try:
        if service_name:
            # Reset specific circuit breaker
            breaker = circuit_breaker_manager._breakers.get(service_name)
            if breaker:
                breaker.reset()
                logger.info("Circuit breaker reset", service=service_name)
                return {
                    "status": "success",
                    "message": f"Circuit breaker for {service_name} reset successfully",
                    "timestamp": datetime.now().isoformat(),
                }
            else:
                return {
                    "status": "error",
                    "message": f"Circuit breaker for {service_name} not found",
                    "available_services": list(
                        circuit_breaker_manager._breakers.keys()
                    ),
                    "timestamp": datetime.now().isoformat(),
                }
        else:
            # Reset all circuit breakers
            circuit_breaker_manager.reset_all()
            logger.info("All circuit breakers reset")
            return {
                "status": "success",
                "message": "All circuit breakers reset successfully",
                "timestamp": datetime.now().isoformat(),
            }

    except Exception as e:
        logger.error(
            "Failed to reset circuit breakers", service=service_name, error=str(e)
        )
        raise HTTPException(
            status_code=500, detail=f"Failed to reset circuit breakers: {str(e)}"
        )
